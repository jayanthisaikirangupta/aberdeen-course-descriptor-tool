"""
Aberdeen Course Descriptor Builder — core logic.

Fetches University of Aberdeen "Catalogue of Courses" pages live, parses the
relevant fields, and compiles a Word (.docx) "Module Descriptors" document that
matches the registry template.

Run from the command line:
    python descriptor_builder.py            # uses config.json
    python descriptor_builder.py myjob.json # uses a different config file

Or import build_all() / build_document() from the Flask app (app.py).
"""

import re
import os
import sys
import json
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "https://www.abdn.ac.uk/registry/courses"
HEADERS = {
    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                   "(KHTML, like Gecko) Chrome/120.0 Safari/537.36"),
    "Accept-Language": "en-GB,en;q=0.9",
}
HERE = os.path.dirname(os.path.abspath(__file__))


class CourseNotFound(Exception):
    """Raised when a course page cannot be located or fetched."""


# --------------------------------------------------------------------------- #
#  URL building                                                               #
# --------------------------------------------------------------------------- #
def prefix_of(code):
    m = re.match(r"[A-Za-z]+", code.strip())
    return m.group(0).upper() if m else ""


def subject_for(code, prefix_map):
    # prefix_map keys are matched case-insensitively
    upper = {k.upper(): v for k, v in prefix_map.items()}
    return upper.get(prefix_of(code))


def candidate_urls(code, level, year, prefix_map, manual_urls=None):
    """Return the list of URLs to try, in priority order.

    Casing matters on the live site, so we try the code/subject as supplied
    *and* a couple of safe fallbacks (UPPER code, lower-case subject path).
    """
    code = code.strip()
    manual_urls = manual_urls or {}
    # explicit manual override wins
    for k, v in manual_urls.items():
        if k.strip().upper() == code.upper() and v.strip():
            return [v.strip()]

    subject = subject_for(code, prefix_map)
    if not subject:
        return []

    root = f"{BASE}/{level}/{year}"
    subjects = [subject]
    if subject.lower() != subject:
        subjects.append(subject.lower())
    codes = [code.upper(), code]

    urls, seen = [], set()
    for s in subjects:
        for c in codes:
            u = f"{root}/{s}/{c}"
            if u not in seen:
                seen.add(u)
                urls.append(u)
    return urls


def fetch_html(code, level, year, prefix_map, manual_urls, session):
    """Fetch the course page HTML, trying URL variants. Returns (url, html)."""
    urls = candidate_urls(code, level, year, prefix_map, manual_urls)
    if not urls:
        raise CourseNotFound(
            f"No subject mapping for prefix '{prefix_of(code)}' (code {code}). "
            f"Add it to prefix_map in config, or supply a manual URL."
        )
    last = None
    for u in urls:
        try:
            r = session.get(u, headers=HEADERS, timeout=25, allow_redirects=True)
        except requests.RequestException as e:
            last = f"network error: {e}"
            continue
        ok_page = (r.status_code == 200
                   and code.upper() in r.text.upper()
                   and "Course Overview" in r.text)
        if ok_page:
            return r.url, r.text
        last = f"HTTP {r.status_code}"
    raise CourseNotFound(
        f"Could not fetch {code} (tried {len(urls)} URL variant(s); last result: {last}). "
        f"The code, level, year or subject mapping may be wrong — supply a manual URL to override."
    )


# --------------------------------------------------------------------------- #
#  HTML parsing                                                               #
# --------------------------------------------------------------------------- #
def _content(soup):
    return (soup.find(id="maincontent") or soup.find("main")
            or soup.body or soup)


def _is_descendant(el, ancestor):
    """True if ``el`` is a descendant of ``ancestor`` (or is ``ancestor``)."""
    node = el
    while node is not None:
        if node is ancestor:
            return True
        node = node.parent
    return False


# Text patterns that indicate we've walked past the page's main content into
# the site's global footer / cookie / accessibility strip. Used as a hard
# stop when scanning forward from a section heading.
_FOOTER_MARKERS = (
    "university of aberdeen is a charity",
    "accessibility statement",
    "freedom of information",
    "privacy statement",
    "king's college aberdeen",
    "©", "©",
)


def _looks_like_footer(text):
    low = (text or "").lower()
    return any(m in low for m in _FOOTER_MARKERS)


def _find_heading(content, keywords):
    for h in content.find_all(["h2", "h3"]):
        t = h.get_text(" ", strip=True).lower()
        if any(k in t for k in keywords):
            return h
    return None


def _section_lines(content, keywords, section_label=""):
    """Collect paragraph / list text under a heading until the next h2/h3.

    Handles the quirks of the Aberdeen catalogue pages:
    - Same paragraph sometimes appears twice in a row (obvious duplicate).
    - A <p> is sometimes followed by a <ul> containing the same sentences
      broken into bullets \u2014 drop the paragraph in that case.
    - A stray 'Course Description:' echo inside the Description section itself.
    """
    head = _find_heading(content, keywords)
    if not head:
        return []
    out = []
    # single set — a bullet <li> and an identical <p> paragraph both count as
    # the same content and should not both appear (Aberdeen pages sometimes
    # duplicate content in both forms).
    seen_norm = set()

    def _norm(s):
        # aggressive: whitespace-collapse, lowercase, strip all non-alphanumerics.
        # catches duplicates that differ only by punctuation/casing/spacing.
        s = re.sub(r"\s+", " ", s or "").strip().lower()
        return re.sub(r"[^a-z0-9 ]", "", s).strip()

    def _covered_by_bullets(text, following):
        bullets = []
        for el in following:
            if el.name in ("h2", "h3", "h4", "h5"):
                break
            if el.name == "li":
                bullets.append(el.get_text(" ", strip=True))
        if not bullets:
            return False
        joined = _norm(" ".join(bullets))
        n = _norm(text)
        if not n:
            return False
        return n == joined or n in joined

    def _is_covered_by_prior(n):
        # already-emitted paragraph fully contains this new one \u2192 skip
        for prev in seen_norm:
            if n and (n == prev or n in prev):
                return True
        return False

    label_norm = _norm(section_label) if section_label else ""
    all_next = head.find_all_next()
    for i, el in enumerate(all_next):
        # stop once we leave the main content container (prevents footer /
        # site-wide navigation from being pulled into a course section)
        if not _is_descendant(el, content):
            break
        if el.name in ("h2", "h3"):
            break
        if el.name == "p":
            # Aberdeen pages sometimes emit malformed <p><p>...</p></p>
            # nesting. Skip a wrapper <p> that itself contains child <p>
            # tags — the inner ones will be visited individually next.
            if el.find("p") is not None:
                continue
            t = el.get_text(" ", strip=True)
            if not t:
                continue
            if _looks_like_footer(t):
                break
            n = _norm(t)
            if not n:
                continue
            if label_norm and n == label_norm:
                continue
            if _is_covered_by_prior(n):
                continue
            if _covered_by_bullets(t, all_next[i + 1:]):
                continue
            seen_norm.add(n)
            out.append(t)
        elif el.name == "li":
            t = el.get_text(" ", strip=True)
            if _looks_like_footer(t):
                break
            n = _norm(t)
            if t and n and n not in seen_norm:
                seen_norm.add(n)
                out.append("\u2022 " + t)
    return out


def _cell_text(cell):
    """Cell text preserving list-item line breaks.

    Cells that hold a <ul>/<ol> (co-ordinators, sometimes assessments) join
    items with '\\n' so multi-person / multi-item values render on separate
    lines. Other cells stay single-line for compact labels/values.
    """
    if cell.find(["ul", "ol"]):
        items = [li.get_text(" ", strip=True)
                 for li in cell.find_all("li")
                 if li.get_text(strip=True)]
        # also pull any stray text that lives OUTSIDE the list within the cell
        list_text = " ".join(items)
        cell_text = cell.get_text(" ", strip=True)
        prefix = cell_text.replace(list_text, "").strip() if items else cell_text
        parts = ([prefix] if prefix else []) + items
        return "\n".join(p for p in parts if p)
    return cell.get_text(" ", strip=True)


def _details_table(content):
    """Parse the 'Course Details' table into a {label: value} dict."""
    d = {}
    for table in content.find_all("table"):
        if "Credit Points" in table.get_text():
            for tr in table.find_all("tr"):
                cells_raw = tr.find_all(["td", "th"])
                cells = [_cell_text(c) for c in cells_raw]
                i = 0
                while i + 1 < len(cells):
                    label, val = cells[i], cells[i + 1]
                    if label:
                        d[label] = val
                    i += 2
            break
    return d


def _credits(details):
    txt = details.get("Credit Points", "")
    cp = (re.search(r"(\d+)\s*credit", txt, re.I) or [None, ""])
    ects = (re.search(r"([\d.]+)\s*ECTS", txt, re.I) or [None, ""])
    cp = cp.group(1) if hasattr(cp, "group") else ""
    ects = ects.group(1) if hasattr(ects, "group") else ""
    return cp, ects


def _coordinators(details):
    for key in ("Co-ordinators", "Coordinators", "Course Co-ordinator",
                "Course Coordinator", "Co-ordinator"):
        if key in details and details[key].strip():
            return details[key].strip()
    return ""


_SKIP_ASSESSMENT_SUBHEADINGS = {
    "learning outcomes", "learning outcome", "look up week numbers",
    "week numbers", "feedback",
}


def _assessment_items(content, keywords):
    """Return (items, free_text) under a section heading.

    ``items`` is a list of (h4-assessment-name, weighting) tuples parsed from the
    newer structured layout. ``free_text`` is any paragraph or bullet text under
    the same heading — used as a fallback for older catalogue years where the
    assessments are described in prose rather than h4 sub-headings.

    h5s on the live site are 'Learning Outcomes' / 'Feedback' subheadings that
    clutter the output, so they're skipped. Filler paragraphs like 'Look up
    Week Numbers' are dropped.
    """
    head = _find_heading(content, keywords)
    if not head:
        return [], ""
    items, texts, seen, seen_items = [], [], set(), set()

    def _norm(s):
        s = re.sub(r"\s+", " ", s or "").strip().lower()
        return re.sub(r"[^a-z0-9 ]", "", s).strip()

    # When we enter a "Feedback" / "Learning outcomes" sub-block (h4 or h5)
    # inside this assessment section, its paragraphs belong to that sub-block
    # (Feedback gets its own rendered section) — skip until we exit it.
    in_skip_subblock = False
    for el in head.find_all_next():
        if not _is_descendant(el, content):
            break
        if el.name in ("h2", "h3"):
            break
        if el.name in ("h4", "h5"):
            sub = el.get_text(" ", strip=True).strip().lower().rstrip(":.")
            in_skip_subblock = sub in _SKIP_ASSESSMENT_SUBHEADINGS
        if el.name == "h4":
            name = el.get_text(" ", strip=True)
            if name.strip().lower() in _SKIP_ASSESSMENT_SUBHEADINGS:
                continue
            weight = None
            tbl = el.find_next("table")
            if tbl:
                m = re.search(r"Weighting\s*([\d.]+)", tbl.get_text(" ", strip=True), re.I)
                if m:
                    weight = m.group(1)
            if name:
                key = (_norm(name), weight)
                if key in seen_items:
                    continue
                seen_items.add(key)
                items.append((name, weight))
        elif el.name in ("p", "li"):
            if in_skip_subblock:
                continue
            t = el.get_text(" ", strip=True)
            if not t:
                continue
            if _looks_like_footer(t):
                break
            low = t.lower()
            if "subject to change" in low:
                continue
            n = _norm(t)
            if not n or n in _SKIP_ASSESSMENT_SUBHEADINGS or n in seen:
                continue
            # skip paragraphs that just restate an already-listed h4 item
            if any(n == item_key or n in item_key or item_key in n
                   for (item_key, _) in seen_items):
                continue
            seen.add(n)
            texts.append(("• " + t) if el.name == "li" else t)

    # Some catalogue pages emit a "combined" summary that concatenates every
    # individual assessment (e.g. "Essay 20% Lab 20% Open book 40%") IN
    # ADDITION to the individual entries — either as an extra h4 or as an
    # extra <p>/<li>. Drop any entry whose text substring-contains 2+ other
    # entries' text; that's the summary form.
    def _drop_summaries(entries, textfn):
        if len(entries) <= 2:
            return entries
        norms = [_norm(textfn(e)) for e in entries]
        out = []
        for i, ni in enumerate(norms):
            covers = sum(1 for j, nj in enumerate(norms)
                         if i != j and nj and nj in ni)
            if covers < 2:
                out.append(entries[i])
        return out

    items = _drop_summaries(items, lambda it: it[0])
    texts = _drop_summaries(texts, lambda t: t)

    return items, "\n".join(texts)


_ALT_RESIT_RE = re.compile(r"alternative\s+resit\s+arrangements?\s*[:\-]?\s*", re.I)


def _find_alt_resit(content):
    """Older catalogue years don't have a Resit Assessments heading — the resit
    info is a bold 'Alternative Resit Arrangements' label inside Summative
    Assessments. Return the text after that label, or '' if not found.
    """
    for tag in content.find_all(["strong", "b"]):
        label = tag.get_text(" ", strip=True)
        if not re.match(r"alternative\s+resit\s+arrangements?", label, re.I):
            continue
        parent = tag.find_parent(["p", "div", "li", "td"]) or tag.parent
        if not parent:
            continue
        full = parent.get_text(" ", strip=True)
        rest = _ALT_RESIT_RE.sub("", full, count=1).strip()
        if rest:
            return rest
    for el in content.find_all(["p", "div", "li", "td"]):
        text = el.get_text(" ", strip=True)
        m = _ALT_RESIT_RE.search(text)
        if m and text[m.end():].strip():
            return text[m.end():].strip()
    return ""


def _fmt_items(items):
    return "\n".join(f"{n}: {w}%" if w else n for n, w in items)


def _feedback_text(content):
    """Collect Feedback sub-section text.

    Aberdeen pages use an h5 'Feedback' heading (usually inside Summative or
    Resit sections) followed by paragraphs describing feedback delivery. We
    walk every 'Feedback' heading found and join the paragraphs beneath it,
    stopping at the next same-or-higher-level heading or the page footer.
    """
    out, seen = [], set()

    def _norm(s):
        s = re.sub(r"\s+", " ", s or "").strip().lower()
        return re.sub(r"[^a-z0-9 ]", "", s).strip()

    for h in content.find_all(["h2", "h3", "h4", "h5"]):
        label = h.get_text(" ", strip=True).lower().strip().rstrip(":.")
        if label != "feedback":
            continue
        stop_levels = {"h2", "h3", "h4", "h5"}
        # allow deeper headings under this one but stop at siblings-or-higher
        this_level = h.name
        for el in h.find_all_next():
            if not _is_descendant(el, content):
                break
            if el.name in stop_levels and el.name <= this_level:
                break
            if el.name not in ("p", "li"):
                continue
            t = el.get_text(" ", strip=True)
            if not t or _looks_like_footer(t):
                if _looks_like_footer(t):
                    break
                continue
            n = _norm(t)
            if not n or n in seen:
                continue
            seen.add(n)
            out.append(("• " + t) if el.name == "li" else t)
    return "\n".join(out)


def _title(content, code):
    h1 = content.find("h1")
    txt = h1.get_text(" ", strip=True) if h1 else code
    # "AN3009: ARCHITECTURE OF LIFE (2024-2025)"
    m = re.match(r"\s*[A-Za-z0-9]+\s*[:\-]\s*(.+?)\s*(?:\(\d{4}-\d{4}\))?\s*$", txt)
    title = m.group(1).strip() if m else txt
    return title.upper()


def parse_course(html, code, url):
    soup = BeautifulSoup(html, "html.parser")
    content = _content(soup)
    details = _details_table(content)
    cp, ects = _credits(details)
    summ_items, summ_text = _assessment_items(content, ["summative assessment"])
    resit_items, resit_text = _assessment_items(content, ["resit"])
    formative_items, formative_text = _assessment_items(content, ["formative assessment"])

    # Older catalogue pages often lack a Resit heading — the info is inline as
    # 'Alternative Resit Arrangements: ...' inside Summative Assessments.
    alt_resit = ""
    if not resit_items and not resit_text:
        alt_resit = _find_alt_resit(content)

    overview = "\n".join(_section_lines(content, ["course overview"], "Course Overview"))
    desc = "\n".join(_section_lines(content, ["course description"], "Course Description"))

    # older catalogue pages use different headings for the same content, or
    # leave the description as a placeholder ("see course page"). Fall back to
    # alternative section names so the descriptor isn't left empty.
    def _too_thin(text):
        stripped = re.sub(r"[^a-z0-9]", "", (text or "").lower())
        return len(stripped) < 40 or stripped in {"seecoursepage", "nocoursedescription"}

    if _too_thin(desc):
        for alt in (["aims and objectives"], ["course aims"],
                    ["syllabus"], ["main learning outcomes"]):
            alt_desc = "\n".join(_section_lines(content, alt, alt[0].title()))
            if not _too_thin(alt_desc):
                desc = alt_desc
                break

    feedback = _feedback_text(content)

    return {
        "id": code.upper(),
        "title": _title(content, code),
        "cp": cp or "—",
        "ects": ects or "—",
        "coord": _coordinators(details) or "See course page",
        "overview": overview or "(no overview found)",
        "desc": desc or "(no description found)",
        "summative": _fmt_items(summ_items) or summ_text or "See course page",
        "formative": _fmt_items(formative_items) or formative_text
                     or "There are no assessments for this course.",
        "resit": _fmt_items(resit_items) or resit_text or alt_resit
                 or "Not specified for this course",
        "feedback": feedback or "Not available for this course",
        "url": url,
    }


# --------------------------------------------------------------------------- #
#  Orchestration                                                              #
# --------------------------------------------------------------------------- #
def _year_groups(config):
    """Return the config's year groups, translating legacy {year, courses}."""
    groups = config.get("years")
    if groups:
        return [g for g in groups if g.get("year") and g.get("courses")]
    if config.get("year") and config.get("courses"):
        return [{"year": config["year"], "courses": config["courses"]}]
    return []


def build_all(config, manual_urls=None, progress=None):
    """Fetch + parse every course. Returns (courses, errors).

    Each course dict carries a ``year`` field indicating the academic year it
    was fetched under. progress: optional callback(code, status, message).
    """
    manual_urls = manual_urls or config.get("manual_urls", {})
    level = config["level"]
    prefix_map = config.get("prefix_map", {})
    session = requests.Session()

    courses, errors = [], []
    for group in _year_groups(config):
        year = group["year"]
        for code in group["courses"]:
            code = code.strip()
            if not code:
                continue
            if progress:
                progress(code, "fetching", "")
            try:
                url, html = fetch_html(code, level, year, prefix_map, manual_urls, session)
                course = parse_course(html, code, url)
                course["year"] = year
                courses.append(course)
                if progress:
                    progress(code, "done", course["title"])
            except CourseNotFound as e:
                errors.append({"code": code.upper(), "year": year, "message": str(e)})
                if progress:
                    progress(code, "error", str(e))
    return courses, errors


# --------------------------------------------------------------------------- #
#  Word document generation                                                   #
# --------------------------------------------------------------------------- #
def _add_page_number(paragraph):
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(instr); run._r.append(f2)


def _set_cell_border(cell, edges=("top", "left", "bottom", "right")):
    """Draw a single black border on the given sides of a table cell.

    Pass ``edges`` to control which edges are drawn — omit e.g. ``"right"``
    to suppress the divider between two adjacent cells in the same row.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        if edge in edges:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tcPr.append(borders)


def _tight(p):
    """Match the tight paragraph spacing used in the sample PDFs.

    NOTE: intentionally does NOT force line_spacing = 1.0. Word's default
    Normal-style spacing (1.15) produces the airy, readable feel of the
    manual reference documents; forcing single spacing made every course
    look cramped.
    """
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


_HYPERLINK_REL = ("http://schemas.openxmlformats.org/officeDocument/2006/"
                  "relationships/hyperlink")


def _add_hyperlink(paragraph, url, text):
    """Append a real clickable hyperlink run to ``paragraph``.

    python-docx has no first-class hyperlink API, so we build the w:hyperlink
    element by hand and register the external relationship on the doc part.
    """
    part = paragraph.part
    r_id = part.relate_to(url, _HYPERLINK_REL, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _label_in(cell, text):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _tight(p)
    r = p.add_run(text); r.bold = True
    return p


def _plain_in(cell, text):
    p = cell.add_paragraph(str(text))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _tight(p)
    return p


def _sublabel_in(cell, text):
    """Bold sub-heading used for 'Summative Assessments', 'Feedback', etc."""
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _tight(p)
    r = p.add_run(str(text))
    r.bold = True
    return p


def _row_cannot_split(row):
    """Force Word to keep this row on a single page instead of splitting it."""
    trPr = row._tr.get_or_add_trPr()
    existing = trPr.find(qn("w:cantSplit"))
    if existing is None:
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)


def _row_keep_with_next(row):
    """Set keep-with-next on every paragraph in the row so Word doesn't
    place this row on one page and the following row on the next.
    """
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.keep_with_next = True


def _text_in(cell, text, bullets=True):
    """Write multi-line text into a cell.

    Lines starting with '•' become real, left-aligned List Bullet list items;
    everything else is a normal left-aligned paragraph. Blank input lines
    become empty paragraphs so the visual spacing follows the source.
    """
    for ln in str(text).split("\n"):
        stripped = ln.lstrip()
        if bullets and stripped.startswith("•"):
            content = stripped[1:].strip()
            p = cell.add_paragraph(content, style="List Bullet")
        else:
            p = cell.add_paragraph(ln)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _tight(p)


def _spacer(cell):
    """A blank paragraph between sections. Full-height (inherits the
    Normal-style font size) so the sections have proper breathing room,
    matching the airy feel of the manual reference documents."""
    return _tight(cell.add_paragraph())


def _course_table(doc, course, page_break_before=False):
    """One course = one bordered table.

    Single-row single-cell OUTER table (atomic; can't be split at a row
    boundary because there's only one row). Inside the cell:
      * a nested 1-row 3-column table acts as the header banner
        (code | title | credit points) — mirrors the manual reference layout,
      * then the body paragraphs (Coordinator, Overview, Description,
        Assessment & Feedback, Feedback).

    ``page_break_before=True`` forces the course to start on a new page
    by setting w:pageBreakBefore on the FIRST paragraph inside the cell —
    this avoids the "phantom empty page" that doc.add_page_break() causes
    (an extra empty paragraph that lands on its own page).
    """
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.allow_autofit = False
    tbl.rows[0].cells[0].width = Inches(6.7)

    body = tbl.rows[0].cells[0]
    # Nuke the empty default paragraph so we start with clean content.
    first_p = body.paragraphs[0]
    first_p._element.getparent().remove(first_p._element)

    # ---- header banner (nested 3-column table) ----
    hdr_tbl = body.add_table(rows=1, cols=3)
    hdr_tbl.autofit = False
    hdr_tbl.allow_autofit = False
    hcells = hdr_tbl.rows[0].cells
    for i, w in enumerate([Inches(1.4), Inches(3.4), Inches(1.9)]):
        hcells[i].width = w

    p = hcells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _tight(p)
    r = p.add_run(course["id"]); r.bold = True
    if page_break_before:
        # Force the whole course onto its own page — set on the very first
        # paragraph of the cell content.
        p.paragraph_format.page_break_before = True

    p = hcells[1].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight(p)
    r = p.add_run(course["title"]); r.bold = True

    p = hcells[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _tight(p)
    r = p.add_run("Credit Points: " + str(course["cp"])); r.bold = True
    p2 = hcells[2].add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _tight(p2)
    r2 = p2.add_run("(ECTS Credits: " + str(course["ects"]) + ")"); r2.bold = True

    # Nested header cells: only a bottom divider so it looks like the
    # original layout (no vertical lines between code/title/credits).
    _set_cell_border(hcells[0], edges=("bottom",))
    _set_cell_border(hcells[1], edges=("bottom",))
    _set_cell_border(hcells[2], edges=("bottom",))
    # Keep the header banner intact (single row, atomic).
    _row_cannot_split(hdr_tbl.rows[0])
    # Glue the header banner to the first body paragraph so the banner
    # never sits alone at the bottom of a page.
    _row_keep_with_next(hdr_tbl.rows[0])

    # ---- body content ----
    _label_in(body, "Course Coordinator(s):")
    _text_in(body, course["coord"], bullets=False)
    _spacer(body)

    _label_in(body, "Course Overview:")
    _text_in(body, course["overview"])
    _spacer(body)

    _label_in(body, "Course Description:")
    _text_in(body, course["desc"])
    _spacer(body)

    _label_in(body, "Assessment & Feedback:")
    _sublabel_in(body, "Summative Assessments")
    _text_in(body, course["summative"])
    _spacer(body)
    _sublabel_in(body, "Formative Assessment")
    _text_in(body, course["formative"], bullets=False)
    _spacer(body)
    _sublabel_in(body, "Resit Assessments")
    _text_in(body, course["resit"])
    _spacer(body)
    _sublabel_in(body, "Feedback")
    _text_in(body, course.get("feedback", "Not available for this course"), bullets=False)

    # Full frame around the whole course. We intentionally DO NOT set
    # cantSplit on the outer row — cantSplit on a row taller than a page
    # makes Word render through the bottom margin (overlapping the footer);
    # letting the row split at paragraph boundaries lets a tall course
    # overflow onto the next page cleanly instead.
    _set_cell_border(body, edges=("top", "left", "bottom", "right"))


def build_document(cover, courses, years, logo_path=None):
    """Build the Module Descriptors document and return a python-docx Document.

    ``years`` may be either a single academic-year string (legacy) or a list of
    ``{year, courses}`` dicts. Courses are grouped by their ``year`` attribute
    (which build_all sets); each group starts on its own page with a bold
    "The following descriptors are correct for the academic year YYYY-YYYY:" line.
    """
    if logo_path is None:
        logo_path = os.path.join(HERE, "assets", "aberdeen-logo.jpeg")

    if isinstance(years, str):
        year_list = [years] if years else []
    else:
        year_list = [g["year"] for g in (years or []) if g.get("year")]

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    # 11pt Calibri matches the readable feel of the manual reference
    # documents. Earlier we shrunk to 10pt to force everything onto single
    # pages, but the result looked cramped; tall courses now spill to a
    # second page (as they did in the original tool output).
    normal.font.size = Pt(11)

    sec = doc.sections[0]
    # Restore the original generous margins that match the manual reference
    # documents' visual polish. Tall courses will flow onto a 2nd page
    # rather than cram everything onto one via tiny margins.
    sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)

    # repeating header logo, with breathing room below so body text
    # doesn't crowd it on any page
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path and os.path.exists(logo_path):
        hp.add_run().add_picture(logo_path, width=Inches(2.3))
    spacer = sec.header.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(12)

    # repeating footer: centered page number, then a left-aligned student block
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(fp)
    for line in (cover.get("student", ""), cover.get("degree", ""),
                 "University of Aberdeen", cover.get("date", "")):
        fpl = footer.add_paragraph(line)
        fpl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fpl.paragraph_format.space_after = Pt(0)

    # ---- cover page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Module Descriptors for the Degree of " + cover.get("degree", ""))
    tr.bold = True

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    intro.add_run(
        "The following descriptors have been taken from the relevant Catalogue of Courses "
        "for the academic year stated during which " + cover.get("student", "") +
        " was a registered student at the University of Aberdeen studying towards the Degree of " +
        cover.get("degree", "") + "."
    )
    p_url = doc.add_paragraph(); p_url.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_url.add_run(
        "The Catalogue of Courses archive can be accessed on the University of Aberdeen "
        "website at the following URL: "
    )
    _add_hyperlink(p_url, "https://www.abdn.ac.uk/registry/courses/",
                   "https://www.abdn.ac.uk/registry/courses/")
    p_url.add_run(".")
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.add_run(
        "At the University of Aberdeen, a credit point is defined as ‘the outcome of learning "
        "achieved by an average student through 10 notional hours of learning time’. Therefore, "
        "courses worth 15 credit points involved 150 notional hours of learning time and courses "
        "worth 30 credit points involved 300 notional hours of learning time."
    )
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.add_run(
        "The University of Aberdeen uses a credit system which is double the ECTS system "
        "(e.g. 5 Aberdeen credit points = 2.5 ECTS credits)."
    )
    doc.add_paragraph()
    doc.add_paragraph("This document was compiled by: " + cover.get("compiled_by", ""))
    doc.add_paragraph("Position: " + cover.get("position", ""))
    doc.add_paragraph("Date: " + cover.get("date", ""))
    doc.add_paragraph("Signature:")
    doc.add_paragraph("Stamp:")

    # ---- course sections, grouped by year ----
    grouped, order = {}, []
    for c in courses:
        y = c.get("year", "")
        if y not in grouped:
            grouped[y] = []
            order.append(y)
        grouped[y].append(c)

    # keep the order the user typed the year groups in, then any strays
    ordered_years = [y for y in year_list if y in grouped]
    ordered_years += [y for y in order if y not in ordered_years]

    first_year = True
    for y in ordered_years:
        intro_year = doc.add_paragraph()
        if y:
            msg = "The following descriptors are correct for the academic year " + y + ":"
        else:
            msg = "The following descriptors are correct as of the year(s) of study below:"
        r = intro_year.add_run(msg)
        r.bold = True
        # Each year intro starts on a new page. Using page_break_before on
        # the paragraph itself (rather than doc.add_page_break()) avoids the
        # "phantom empty page" caused by a stray break paragraph.
        intro_year.paragraph_format.page_break_before = True
        # Glue the year intro to the first course so they share a page.
        intro_year.paragraph_format.keep_with_next = True

        for idx, c in enumerate(grouped[y]):
            # Let courses flow naturally within the year — matches the
            # manual reference documents where a short course may share a
            # page with the tail of another. The atomic single-row table
            # (created in _course_table) still keeps each course's header
            # glued to its own body, so no header/body split occurs.
            _course_table(doc, c, page_break_before=False)
        first_year = False

    return doc


def safe_name(s):
    return re.sub(r"[^A-Za-z0-9]+", "_", s or "Descriptors").strip("_") or "Descriptors"


# --------------------------------------------------------------------------- #
#  CLI                                                                        #
# --------------------------------------------------------------------------- #
def main():
    cfg_path = sys.argv[1] if len(sys.argv) > 1 else os.path.join(HERE, "config.json")
    with open(cfg_path, "r", encoding="utf-8") as f:
        config = json.load(f)

    def show(code, status, msg):
        icon = {"fetching": "...", "done": " ok", "error": "ERR"}.get(status, "   ")
        print(f"  [{icon}] {code}{'  ' + msg if msg else ''}")

    groups = _year_groups(config)
    total = sum(len(g["courses"]) for g in groups)
    years_txt = ", ".join(g["year"] for g in groups) or "no years configured"
    print(f"Fetching {total} course(s) ({config['level']} — {years_txt})...")
    courses, errors = build_all(config, progress=show)

    if not courses:
        print("\nNo courses fetched successfully. Nothing to write.")
        if errors:
            for e in errors:
                print(f"  - {e['code']}: {e['message']}")
        sys.exit(1)

    out_dir = os.path.join(HERE, "output")
    os.makedirs(out_dir, exist_ok=True)
    cover = config["cover"]
    doc = build_document(cover, courses, groups)
    out_path = os.path.join(out_dir, "Module_Descriptors_" + safe_name(cover.get("student")) + ".docx")
    doc.save(out_path)

    print(f"\nWrote {len(courses)} course(s) to:\n  {out_path}")
    if errors:
        print(f"\n{len(errors)} course(s) FAILED (not in the document):")
        for e in errors:
            print(f"  - {e['code']}: {e['message']}")
        print("\nAdd a manual URL for these in config.json -> manual_urls, then re-run.")


if __name__ == "__main__":
    main()
